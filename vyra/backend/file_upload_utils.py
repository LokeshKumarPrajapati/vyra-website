"""
File Upload Utilities for VYRA
Handles parsing and content extraction from various file formats:
- TXT, CSV, JSON, Excel (.xlsx, .xls), Word (.docx), PDF
"""

import json
from typing import Dict, Any

# File parsing libraries
import pandas as pd
from docx import Document
import PyPDF2

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB in bytes

SUPPORTED_FORMATS = {
    'txt': 'text/plain',
    'csv': 'text/csv',
    'json': 'application/json',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'xls': 'application/vnd.ms-excel',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'pdf': 'application/pdf'
}


def parse_txt(file_path: str) -> Dict[str, Any]:
    """Parse text file and return content"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return {
            'success': True,
            'content': content,
            'summary': f"Text file with {len(content)} characters, {len(content.splitlines())} lines",
            'type': 'text'
        }
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as f:
            content = f.read()

        return {
            'success': True,
            'content': content,
            'summary': f"Text file with {len(content)} characters (latin-1 encoding)",
            'type': 'text'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'type': 'text'
        }


def parse_csv(file_path: str) -> Dict[str, Any]:
    """Parse CSV file and return structured data"""
    try:
        df = pd.read_csv(file_path)

        # Convert to readable format
        rows, cols = df.shape
        content = f"CSV Data ({rows} rows × {cols} columns)\n\n"
        content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
        content += "Preview (first 10 rows):\n"
        content += df.head(10).to_string()

        return {
            'success': True,
            'content': content,
            'data': df.to_dict('records'),  # Full data as list of dicts
            'summary': f"CSV file with {rows} rows and {cols} columns",
            'type': 'csv'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'type': 'csv'
        }


def parse_json(file_path: str) -> Dict[str, Any]:
    """Parse JSON file and return structured data"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Create readable representation
        content = json.dumps(data, indent=2, ensure_ascii=False)

        return {
            'success': True,
            'content': content,
            'data': data,
            'summary': f"JSON file with {len(content)} characters",
            'type': 'json'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'type': 'json'
        }


def parse_excel(file_path: str) -> Dict[str, Any]:
    """Parse Excel file (.xlsx, .xls) and return structured data"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheets_data = {}
        content = f"Excel Workbook with {len(excel_file.sheet_names)} sheet(s)\n\n"

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            rows, cols = df.shape

            content += f"Sheet: '{sheet_name}' ({rows} rows × {cols} columns)\n"
            content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            content += "Preview (first 5 rows):\n"
            content += df.head(5).to_string()
            content += "\n\n" + "="*50 + "\n\n"

            sheets_data[sheet_name] = df.to_dict('records')

        return {
            'success': True,
            'content': content,
            'data': sheets_data,
            'summary': f"Excel file with {len(excel_file.sheet_names)} sheet(s)",
            'type': 'excel'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'type': 'excel'
        }


def parse_word(file_path: str) -> Dict[str, Any]:
    """Parse Word document (.docx) and extract text"""
    try:
        doc = Document(file_path)

        # Extract all paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)

        # Extract tables if any
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_content.append(table_data)

        if tables_content:
            content += f"\n\n[Document contains {len(tables_content)} table(s)]"

        return {
            'success': True,
            'content': content,
            'summary': f"Word document with {len(paragraphs)} paragraphs, {len(content)} characters",
            'type': 'word',
            'tables': tables_content if tables_content else None
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'type': 'word'
        }


def parse_pdf(file_path: str) -> Dict[str, Any]:
    """Parse PDF file and extract text"""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            # Extract text from all pages
            content = ""
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                content += f"\n--- Page {page_num + 1} ---\n"
                content += page_text

            return {
                'success': True,
                'content': content,
                'summary': f"PDF document with {num_pages} page(s), {len(content)} characters",
                'type': 'pdf',
                'pages': num_pages
            }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'type': 'pdf'
        }


def parse_file(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Parse file based on its type

    Args:
        file_path: Path to the file
        file_type: File extension (txt, csv, json, xlsx, xls, docx, pdf)

    Returns:
        Dictionary with parsing results
    """
    parsers = {
        'txt': parse_txt,
        'csv': parse_csv,
        'json': parse_json,
        'xlsx': parse_excel,
        'xls': parse_excel,
        'docx': parse_word,
        'pdf': parse_pdf
    }

    if file_type not in parsers:
        return {
            'success': False,
            'error': f'Unsupported file type: {file_type}',
            'type': file_type
        }

    return parsers[file_type](file_path)


def validate_file(filename: str, file_size: int) -> Dict[str, Any]:
    """
    Validate file before upload

    Args:
        filename: Name of the file
        file_size: Size in bytes

    Returns:
        Dictionary with validation results
    """
    # Check file extension
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    if ext not in SUPPORTED_FORMATS:
        return {
            'valid': False,
            'error': f'Unsupported file format. Supported: {", ".join(SUPPORTED_FORMATS.keys())}'
        }

    # Check file size
    if file_size > MAX_FILE_SIZE:
        max_mb = MAX_FILE_SIZE / (1024 * 1024)
        actual_mb = file_size / (1024 * 1024)
        return {
            'valid': False,
            'error': f'File too large ({actual_mb:.1f}MB). Maximum size: {max_mb:.0f}MB'
        }

    return {
        'valid': True,
        'extension': ext,
        'mime_type': SUPPORTED_FORMATS[ext]
    }


def get_file_summary(parsed_result: Dict[str, Any], filename: str) -> str:
    """
    Generate a summary of the parsed file for AI context

    Args:
        parsed_result: Result from parse_file()
        filename: Name of the file

    Returns:
        Formatted summary string for AI
    """
    if not parsed_result.get('success'):
        return f"Failed to parse {filename}: {parsed_result.get('error')}"

    summary = f"📎 File Upload: {filename}\n"
    summary += f"Type: {parsed_result.get('type', 'unknown').upper()}\n"
    summary += f"Summary: {parsed_result.get('summary', 'No summary available')}\n\n"

    # Truncate content if too long (keep first 5000 chars for AI context)
    content = parsed_result.get('content', '')
    if len(content) > 5000:
        summary += "Content (truncated):\n"
        summary += content[:5000] + \
            "\n\n[... content truncated, full document available ...]"
    else:
        summary += "Content:\n"
        summary += content

    return summary
